import os
import tempfile
import fitz  # PyMuPDF
import docx
from PIL import Image
from parsers.ocr_engine import run_ocr
import re

def extract_text_from_pdf(pdf_path, logger=None):
    """
    Extracts text from a PDF file.
    Performs layout analysis on the first page to identify the largest font size text
    (highly likely to be the candidate's name).
    If the text is empty or very short, falls back to OCR.
    """
    if logger:
        logger.info(f"Extracting text from PDF: {os.path.basename(pdf_path)}")
        
    doc = fitz.open(pdf_path)
    full_text = []
    largest_text = ""
    largest_font_size = 0.0
    
    # Simple regex to filter out email, phone, web, etc. when looking for a name
    non_name_pattern = re.compile(
        r"@|\+?\d{8,}|\b(email|phone|mobile|tel|address|curriculum|vitae|resume|page|contact|http|www)\b", 
        re.IGNORECASE
    )
    
    for page_idx, page in enumerate(doc):
        # Extract text blocks
        page_text = page.get_text("text")
        full_text.append(page_text)
        
        # On the first page, analyze block layout to find the candidate's name (largest text)
        if page_idx == 0:
            try:
                blocks = page.get_text("dict")["blocks"]
                for b in blocks:
                    if "lines" in b:
                        for l in b["lines"]:
                            for s in l["spans"]:
                                text_val = s["text"].strip()
                                font_size = s["size"]
                                # Filter out short chunks, numbers, or elements containing emails/phones
                                if (len(text_val) > 2 and 
                                    len(text_val) < 60 and 
                                    font_size > largest_font_size and 
                                    not non_name_pattern.search(text_val)):
                                    # Basic check: should have letters
                                    if any(c.isalpha() for c in text_val):
                                        largest_text = text_val
                                        largest_font_size = font_size
            except Exception as e:
                if logger:
                    logger.warning(f"Failed to perform layout analysis on first page: {e}")
                    
    combined_text = "\n".join(full_text)
    
    # If PDF is scanned (very little text extracted), run OCR page by page
    if len(combined_text.strip()) < 100:
        if logger:
            logger.info("Extracted text is empty or too short. PDF is likely scanned. Initiating page-by-page OCR...")
        ocr_pages_text = []
        
        # Create a temporary directory to save page images
        with tempfile.TemporaryDirectory() as temp_dir:
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                # Render page to image (zoom factor of 2.0 for better OCR resolution)
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                image_path = os.path.join(temp_dir, f"page_{page_idx}.png")
                pix.save(image_path)
                
                # Run OCR on the page image
                page_ocr_text = run_ocr(image_path, logger)
                ocr_pages_text.append(page_ocr_text)
                
        combined_text = "\n".join(ocr_pages_text)
        
    doc.close()
    
    # Clean up name string
    largest_text = clean_extracted_name(largest_text)
    
    return {
        "text": combined_text,
        "potential_name": largest_text if largest_text else None
    }

def extract_text_from_docx(docx_path, logger=None):
    """
    Extracts text from a DOCX file, including text in tables.
    """
    if logger:
        logger.info(f"Extracting text from DOCX: {os.path.basename(docx_path)}")
        
    doc = docx.Document(docx_path)
    full_text = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text)
            
    # Extract from tables (very common in academic CVs)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Deduplicate cell text if it spans multiple columns/rows
                cell_text = cell.text.strip()
                if cell_text and cell_text not in row_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))
                
    combined_text = "\n".join(full_text)
    
    return {
        "text": combined_text,
        "potential_name": None # DOCX doesn't easily expose visual font layout similarly, so fallback to text parsing
    }

def extract_text_from_image(image_path, logger=None):
    """
    Extracts text from an image file using OCR.
    """
    ocr_text = run_ocr(image_path, logger)
    return {
        "text": ocr_text,
        "potential_name": None
    }

def clean_extracted_name(name_str):
    """Helper to sanitize and clean extracted name candidates."""
    if not name_str:
        return ""
    # Remove bullet points, tabs, excessive whitespace, punctuation at ends
    name_str = re.sub(r"^[\s\-\*•\d\.\,\(\)]+", "", name_str)
    name_str = re.sub(r"[\s\-\*•\d\.\,\(\)]+$", "", name_str)
    name_str = re.sub(r"\s+", " ", name_str).strip()
    
    # If name is still generic, return empty
    if name_str.lower() in ["curriculum vitae", "resume", "biodata", "bio-data", "cv", "portfolio"]:
        return ""
    return name_str

def parse_file(file_path, logger=None):
    """
    Detects file type and triggers the correct extraction pipeline.
    """
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        result = extract_text_from_pdf(file_path, logger)
        file_type = "PDF"
    elif ext in [".docx", ".doc"]:
        # doc is handled if python-docx can open it, but python-docx only supports docx natively.
        # Users uploading doc should be warned, or we attempt to parse. Let's label as DOCX.
        result = extract_text_from_docx(file_path, logger)
        file_type = "DOCX"
    elif ext in [".png", ".jpg", ".jpeg"]:
        result = extract_text_from_image(file_path, logger)
        file_type = "Image"
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return {
        "filename": os.path.basename(file_path),
        "file_type": file_type,
        "raw_text": result["text"],
        "potential_name": result["potential_name"]
    }
